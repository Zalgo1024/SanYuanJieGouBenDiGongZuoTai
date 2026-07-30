"""
Word 渲染器 — 将解析后的报告数据渲染为 .docx 文件。

排版格式参考：韩庆云伪造8.5亿存款证明中标35亿元水利工程案例分析报告
- 字体：Times New Roman 全线
- 正文：12pt，首行缩进 0.85cm，颜色 #333333
- 标题：16pt 深蓝 #1B3A5C 加粗
- 引用块：单格浅灰背景表格（F5F7FA）+ ▌ 前缀
- 结论箭头（⇒）：单格浅红背景表格（FFF5F5）+ ⇒ 前缀
- 封面标题：26pt 深蓝 #1B3A5C
- 封面版权：10.5pt 灰色 #7F8C8D
- 页边距：上3.70 / 下3.50 / 左2.80 / 右2.60 cm

依赖: python-docx (v1.2+)
"""

import os
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import Config
from parser import Block, ParsedReport, Section

# ── 模式注册表（数据驱动，取代硬编码 if/elif）──
# 新增任意模式 = 往表里加一项，渲染器零改动。
# 单模式按 canonical 序（向后兼容「别动老东西」）；多模式按作者源序（灵活组合）。
MODULES = {
    "policy": {
        "sections": ["fact_summary", "framework", "policy_portrait",
                     "policy_weight", "analysis_body", "conclusion", "appendix"],
        "sentinels": ["policy_portrait", "policy_weight"],
        "label": "政策",
    },
    "event": {
        "sections": ["fact_summary", "framework", "case_portrait",
                     "case_flows", "analysis_body", "case_dynamics",
                     "conclusion", "appendix"],
        "sentinels": ["case_portrait", "case_flows", "case_dynamics"],
        "label": "事件/案例",
    },
    "org": {
        "sections": ["org_portrait", "org_structure", "org_survival",
                     "org_reproduction", "org_interest_network", "org_reverse",
                     "org_transformation", "conclusion", "appendix"],
        "sentinels": ["org_portrait", "org_structure", "org_survival",
                      "org_reproduction", "org_interest_network", "org_reverse",
                      "org_transformation"],
        "label": "组织",
    },
    "opinion": {
        "sections": ["opinion_event", "opinion_actors", "opinion_narrative",
                     "opinion_trilife", "opinion_reverse", "opinion_evolution",
                     "conclusion", "appendix"],
        "sentinels": ["opinion_event", "opinion_actors", "opinion_narrative",
                      "opinion_trilife", "opinion_reverse", "opinion_evolution"],
        "label": "舆情",
    },
}



# ── 正则 ─────────────────────────────────────────────────────

_RE_BOLD = re.compile(r"(?<!\*)\*\*(.+?)\*\*(?!\*)")      # **粗体**（避开 ***）
_RE_ARROW = re.compile(r"^([→⇒])\s*(.*)")                    # → / ⇒ 开头的行
_RE_HEADING = re.compile(r"^(#{1,4})\s+(.*)")                # # / ## / ### / #### 标题
_RE_CONCLUSION_LABEL = re.compile(r"^\*\*(.+?)：\*\*(.*)")   # **汇流段**：...
_RE_LONG_DASH = re.compile(r"—{2,}")                          # 两个及以上连续破折号

# 颜色常量
_COVER_TITLE_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
_COVER_MUTED_COLOR = RGBColor(0x7F, 0x8C, 0x8D)
_H1_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
_H2_COLOR = RGBColor(0x4F, 0x81, 0xBD)
_BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
_ARROW_COLOR = RGBColor(0xC0, 0x39, 0x2B)
_QUOTE_BG = "F5F7FA"
_CONCLUSION_BG = "FFF5F5"

_CHINESE_DIGITS = "零一二三四五六七八九"
_SECTION_NUMBER_PREFIX = re.compile(
    r"^\s*(?:第\s*)?(?:[一二三四五六七八九十百千万零〇两]+|\d+)\s*[、.．)]\s*"
)


def _chinese_number(number: int) -> str:
    """将正整数转为适合章节标题的中文序号（如 1 -> 一、11 -> 十一）。"""
    if number <= 0:
        raise ValueError("章节序号必须是正整数")
    if number < 10:
        return _CHINESE_DIGITS[number]
    if number < 20:
        return "十" + (_CHINESE_DIGITS[number % 10] if number % 10 else "")
    if number < 100:
        tens, ones = divmod(number, 10)
        return f"{_CHINESE_DIGITS[tens]}十" + (_CHINESE_DIGITS[ones] if ones else "")
    if number < 1000:
        hundreds, remainder = divmod(number, 100)
        tens, ones = divmod(remainder, 10)
        result = f"{_CHINESE_DIGITS[hundreds]}百"
        if remainder == 0:
            return result
        if remainder < 10:
            return result + "零" + _CHINESE_DIGITS[ones]
        return result + ("零" if tens == 0 else _CHINESE_DIGITS[tens] + "十") + (
            _CHINESE_DIGITS[ones] if ones else ""
        )
    raise ValueError("章节序号暂不支持超过 999")


def format_section_title(title: str, number: int) -> str:
    """为章节标题统一补齐中文序号，并避免重复编号。"""
    clean_title = _SECTION_NUMBER_PREFIX.sub("", title.strip())
    return f"{_chinese_number(number)}、{clean_title}"


def _compress_dashes(text: str) -> str:
    """将连续 2 个以上的破折号压缩为 1 个（仅用于渲染，不改原文）。

    双重全角破折号（——）是 AI 写作的典型痕迹，压缩为单破折号（—）
    可显著降低「机器味」，且不改变语义。
    """
    return _RE_LONG_DASH.sub("—", text)


def _hex_color(hex_str: str) -> RGBColor:
    """'1B3A5C' → RGBColor(0x1B, 0x3A, 0x5C)"""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── 主入口 ───────────────────────────────────────────────────

def render_docx(
    report: ParsedReport,
    output_path: str,
    config: Optional[Config] = None,
    output_folder: Optional[str] = None,
    diagram_collector: Optional[list] = None,
    tone: str = "neutral",
) -> str:
    if config is None:
        from config import load_config
        config = load_config()

    doc = Document()

    _setup_page(doc, config)
    _render_cover(doc, report, config, tone)
    _render_toc(doc, report, config)  # ← 目录页

    # 进入正文节：添加页眉 + 页码
    _setup_header_footer(doc, report, config)

    _diagram_counter = [0]  # 可变列表，供闭包修改

    # ── 源序组合路由（取代 if/elif 单选）──
    # 检测出现过的模式（哨兵 = 各模式 sentinels 命中 _cid_index）
    present = [m for m in MODULES
               if any(s in report._cid_index for s in MODULES[m]["sentinels"])]
    mode = None
    base = []

    if not present:
        # 无哨兵命中（纯自定义结构，如最终改写稿）：按源序全量渲染，避免整段丢失
        ordered = list(report.section_seq)
    elif len(present) == 1:
        # 单模式：按 canonical 序渲染 = 与历史输出逐字一致（别动老东西）
        mode = present[0]
        base = MODULES[mode]["sections"]
        ordered = []
        for cid in base:
            if cid in report._cid_index or cid in report.sections:
                sec = next((s for s in report.section_seq if s.cid == cid), None)
                if sec is None:
                    # 兼容测试与外部调用直接构造 ParsedReport 的旧接口。
                    sec = report.sections.get(cid)
                if sec is not None:
                    if sec.cid is None:
                        sec.cid = cid
                    ordered.append(sec)
    else:
        # 多模式：按作者源序渲染（灵活多变：事件+舆情+组织+政策任意子集）
        ordered = list(report.section_seq)

    for section_number, sec in enumerate(ordered, start=1):
        if mode == "_base":
            # 共享基础模式保留事实摘要/分析框架/正文等 canonical 槽位。
            section_number = base.index(sec.cid) + 1
        _render_section(doc, sec, config, sec.cid,
                        section_number=section_number,
                        output_folder=output_folder,
                        diagram_collector=diagram_collector,
                        diagram_counter=_diagram_counter)

    doc.save(output_path)

    # 设置自动更新域（TOC 在打开 Word 时自动填充页码）
    _set_auto_update_fields(output_path)

    return output_path


def _set_auto_update_fields(docx_path: str) -> None:
    """设置文档属性：打开时自动更新域（使 TOC 自动填充页码）。"""
    import zipfile
    from lxml import etree

    # 读取 docx 作为 ZIP，修改 settings.xml
    import tempfile
    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    root = etree.fromstring(data)
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    # 如果已有 updateFields 则修改值，否则新建
                    uf = root.find(".//w:updateFields", ns)
                    if uf is None:
                        uf = etree.SubElement(root, f'{{{ns["w"]}}}updateFields')
                        uf.set(f'{{{ns["w"]}}}val', "1")
                    else:
                        uf.set(f'{{{ns["w"]}}}val', "1")
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                          standalone=True)
                zout.writestr(item, data)

    # 落位：覆盖写 docx。
    # shutil.move 在 Windows 上若目标已存在会回退为 copyfile(覆盖 docx)+unlink(tmp)；
    # 沙箱对 os.unlink 有安全拦截（tmp 会被送回收站且可能 fail-closed）。这里改为显式
    # copyfile 覆盖 docx，并忽略 tmp 的清理异常（tmp 为临时文件，残留无害），确保
    # docx 内容正确落位且不因临时文件清理失败而中断流程。真机环境无此 shim，行为一致。
    import os
    import shutil
    try:
        shutil.move(tmp_path, docx_path)
    except Exception:  # noqa: BLE001
        shutil.copyfile(tmp_path, docx_path)
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ── 页面设置 ──────────────────────────────────────────────

def _setup_page(doc: Document, config: Config) -> None:
    typo = config.typography
    page = typo.get("page", {})

    section = doc.sections[0]
    section.top_margin = Cm(page.get("margin_top_mm", 37.0) / 10)
    section.bottom_margin = Cm(page.get("margin_bottom_mm", 35.0) / 10)
    section.left_margin = Cm(page.get("margin_left_mm", 28.0) / 10)
    section.right_margin = Cm(page.get("margin_right_mm", 26.0) / 10)

    # 设置 Normal 样式默认字体 + 段间距
    _set_font(doc.styles["Normal"], "body", config)
    spacing_pt = typo.get("paragraph_spacing_pt", 6)
    doc.styles["Normal"].paragraph_format.space_after = Pt(spacing_pt)

    # 清除初始节的页眉页脚
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ""
    for p in section.footer.paragraphs:
        p.text = ""

    # ── 配置标题样式 ──
    _setup_heading_style(doc.styles["Heading 1"], typo, "heading1",
                         _H1_COLOR, typo.get("font_size", {}).get("heading1", 16))
    _setup_heading_style(doc.styles["Heading 2"], typo, "heading2",
                         _H2_COLOR, typo.get("font_size", {}).get("heading2", 13))
    _setup_heading_style(doc.styles["Heading 3"], typo, "heading2",
                         _H2_COLOR, typo.get("font_size", {}).get("heading2", 12))


def _setup_heading_style(style, typo: dict, font_key: str,
                         color: RGBColor, size_pt: int) -> None:
    """配置标题样式的字体、颜色、间距。"""
    fname = typo.get("font_family", {}).get(font_key, "Times New Roman")
    style.font.name = fname
    style.font.size = Pt(size_pt)
    style.font.bold = True
    style.font.color.rgb = color
    style.element.rPr.rFonts.set(qn("w:eastAsia"), fname)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)


def _add_section_break(doc: Document) -> None:
    new_section = doc.add_section()
    prev = doc.sections[-2]
    new_section.top_margin = prev.top_margin
    new_section.bottom_margin = prev.bottom_margin
    new_section.left_margin = prev.left_margin
    new_section.right_margin = prev.right_margin


# ── 字体工具 ──────────────────────────────────────────────

def _set_font(style_or_run, font_key: str, config: Config) -> None:
    """为 style 或 run 设置字体。"""
    typo = config.typography
    fname = typo.get("font_family", {}).get(font_key, "Times New Roman")
    fsize = typo.get("font_size", {}).get(font_key, 12)
    style_or_run.font.name = fname
    style_or_run.font.size = Pt(fsize)
    # eastAsia 字体也设成一样的（避免中文 fallback 到宋体）
    try:
        style_or_run.element.rPr.rFonts.set(qn("w:eastAsia"), fname)
    except Exception:
        pass


def _add_run(para, text: str, font_key: str, config: Config,
             bold: bool = False, color: RGBColor = _BODY_COLOR, size_pt: Optional[float] = None):
    """快捷添加 run。"""
    run = para.add_run(text)
    _set_font(run, font_key, config)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color
    return run


# ── 封面页 ────────────────────────────────────────────────

def _render_cover(doc: Document, report: ParsedReport, config: Config, tone: str = "neutral") -> None:
    typo = config.typography
    before = typo.get("cover_spacing_lines_before_title", 6)
    after = typo.get("cover_spacing_lines_after_title", 3)

    for _ in range(before):
        doc.add_paragraph("")

    # 报告标题 — 26pt 深蓝 加粗 居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, report.title or "未命名报告", "title", config,
             bold=True, color=_COVER_TITLE_COLOR)

    # 空行
    for _ in range(after):
        doc.add_paragraph("")

    # 版权声明 — 10.5pt 灰色 居中
    cr_text = f"分析框架：{config.theory.get('copyright', '三元结构理论 © 2026, CC BY-NC-SA 4.0')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, cr_text, "footnote", config, color=_COVER_MUTED_COLOR)

    # 生成日期
    import datetime
    now_str = datetime.datetime.now().strftime("%Y-%m-%d")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"报告撰写日期：{now_str}", "footnote", config, color=_COVER_MUTED_COLOR)

    # 分析基调标注 — 与版权/日期同款灰色居中，让「写之前定调」在成品上可见
    _TONE_LABELS = {"neutral": "客观中立", "provocative": "煽动性"}
    _tone_key = tone if tone in _TONE_LABELS else "neutral"
    tone_label = _TONE_LABELS[_tone_key]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"分析基调：{tone_label}", "footnote", config, color=_COVER_MUTED_COLOR)
    # 记录到文档属性，便于程序 / 后续自动化读取
    try:
        doc.core_properties.comments = f"tone={_tone_key}"
    except Exception:
        pass

    _add_section_break(doc)
    # 封面节：清除页眉页脚
    sec = doc.sections[-1]
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

def _render_toc(doc: Document, report: ParsedReport, config: Config) -> None:
    """封面之后插入目录页。

    使用 Word TOC 域代码：自动从 Heading 1/2 样式捕获标题和页码。
    在 Word 中打开后按 Ctrl+A → F9 更新域即可生成带页码的目录。
    """
    # 目录标题
    p = doc.add_paragraph("目  录", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)

    # 插入 TOC 域代码
    # TOC \o "1-2" — 从 Heading 1-2 样式构建
    # \h — 超链接（可点击跳转）
    # \z — 隐藏制表符前导符（保留页码）
    # \u — 使用样式的大纲级别
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()

    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-2" \\h \\z \\u '
        f'</w:instrText>'
    )
    fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)

    # 回退文本：域未更新时的占位显示
    fallback_run = toc_para.add_run("（目录将在 Word 中更新域后生成）")
    _set_font(fallback_run, "body", config)
    fallback_run.font.size = Pt(11)
    fallback_run.font.color.rgb = _COVER_MUTED_COLOR
    fallback_run.italic = True

    run._element.append(fld_end)

    # 分页进入正文
    doc.add_paragraph("")
    doc.add_page_break()


# ── 页眉 + 页码 ───────────────────────────────────────────

def _setup_header_footer(doc: Document, report: ParsedReport, config: Config) -> None:
    """为最新节（正文节）设置页眉和页码。"""
    section = doc.sections[-1]

    # ── 页眉：报告标题 ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = report.title or "三元结构理论案例分析报告"
    _add_run(hp, title_text, "footnote", config, color=_COVER_MUTED_COLOR)

    # ── 页码：第X页/共Y页（居中） ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "第"
    _add_run(fp, "第 ", "footnote", config, color=_COVER_MUTED_COLOR)

    # { PAGE } 域
    run_p = fp.add_run()
    run_p.font.size = Pt(config.typography.get("font_size", {}).get("footnote", 10.5))
    run_p.font.color.rgb = _COVER_MUTED_COLOR
    _add_field_code(run_p, " PAGE ")

    # "页 / 共"
    _add_run(fp, " 页 / 共 ", "footnote", config, color=_COVER_MUTED_COLOR)

    # { NUMPAGES } 域
    run_n = fp.add_run()
    run_n.font.size = Pt(config.typography.get("font_size", {}).get("footnote", 10.5))
    run_n.font.color.rgb = _COVER_MUTED_COLOR
    _add_field_code(run_n, " NUMPAGES ")

    # "页"
    _add_run(fp, " 页", "footnote", config, color=_COVER_MUTED_COLOR)


def _add_field_code(run, instr_text: str) -> None:
    """向 run 中添加 Word 域代码（如 PAGE / NUMPAGES）。"""
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{instr_text}</w:instrText>')
    fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


# ── 章节渲染 ──────────────────────────────────────────────

def _render_section(doc: Document, section: Section, config: Config, sec_id: str,
                    section_number: Optional[int] = None,
                    output_folder: Optional[str] = None,
                    diagram_collector: Optional[list] = None,
                    diagram_counter: Optional[list] = None) -> None:
    title = (format_section_title(section.title, section_number)
             if section_number is not None else section.title)
    doc.add_paragraph(title, style="Heading 1")

    for block in section.blocks:
        _render_block(doc, block, config, sec_id,
                      output_folder=output_folder,
                      diagram_collector=diagram_collector,
                      diagram_counter=diagram_counter)


def _render_block(doc: Document, block: Block, config: Config, sec_id: str,
                  output_folder: Optional[str] = None,
                  diagram_collector: Optional[list] = None,
                  diagram_counter: Optional[list] = None) -> None:
    if block.type == "blank":
        return

    elif block.type == "paragraph":
        text = block.text.strip()
        if not text:
            return

        # ### / #### 子标题
        m = _RE_HEADING.match(text)
        if m:
            level = len(m.group(1))
            title_text = m.group(2).strip()
            if level <= 2:
                doc.add_paragraph(title_text, style="Heading 1")
            elif level == 3:
                doc.add_paragraph(title_text, style="Heading 2")
            else:
                doc.add_paragraph(title_text, style="Heading 3")
            return

        # **汇流段**： / **核心判断**： 等加粗标签行
        m_label = _RE_CONCLUSION_LABEL.match(text)
        if m_label:
            para = doc.add_paragraph()
            _add_run(para, m_label.group(1) + "：", "body", config,
                     bold=True, color=_BODY_COLOR)
            _add_rich_text(para, m_label.group(2).strip(), config,
                           segments=block.segments)
            return

        # ⇒ 开头的行 → 浅红背景单格表格（引用报告格式）
        if text.startswith("⇒"):
            _render_single_cell_table(doc, text, _CONCLUSION_BG, "⇒", config)
            return

        # → 开头的行 → 保持为正文（箭头红色）
        if text.startswith("→"):
            para = doc.add_paragraph()
            _add_rich_text(para, text, config, segments=block.segments)
            _add_first_line_indent(para, config)
            return

        # 普通正文段落
        para = doc.add_paragraph()
        _add_rich_text(para, text, config, segments=block.segments)
        _add_first_line_indent(para, config)

    elif block.type == "quote":
        # 引用块 → 浅灰背景单格表格 + ▌ 前缀
        _render_single_cell_table(doc, f"▌ {block.text}", _QUOTE_BG, None, config)

    elif block.type == "table":
        doc.add_paragraph("")
        _render_multi_row_table(doc, block, config)
        doc.add_paragraph("")

    elif block.type == "list":
        # 附录中的来源清单改用数字编号，便于核对与引用；其余保持项目符号。
        list_style = "List Number" if sec_id == "appendix" else "List Bullet"
        for idx, item_text in enumerate(block.items or []):
            para = doc.add_paragraph(style=list_style)
            # 提取该列表项的链接分段（如果有）
            item_segs = None
            if block.segments and idx < len(block.segments):
                item_segs = block.segments[idx]
            _add_rich_text(para, item_text, config, segments=item_segs)

    elif block.type == "diagram":
        diag = _render_diagram(doc, block, config, output_folder=output_folder,
                               diagram_counter=diagram_counter)
        if diag and diagram_collector is not None:
            diagram_collector.append(diag)

    elif block.type == "subheading":
        # 子维度标题（###），渲染为 Heading 2 样式
        doc.add_paragraph(block.text, style="Heading 2")

    elif block.type == "heading":
        # 回退：未识别类别的标题块，按正文段落处理避免静默丢失
        para = doc.add_paragraph()
        _add_rich_text(para, block.text, config)
        _add_first_line_indent(para, config)


# ── 首行缩进 ──────────────────────────────────────────────

def _add_first_line_indent(para, config: Config) -> None:
    indent_cm = config.typography.get("first_line_indent_cm", 0.85)
    para.paragraph_format.first_line_indent = Cm(indent_cm)


# ── 单格表格（引用 / 结论） ──────────────────────────────

def _render_single_cell_table(doc: Document, text: str, bg_fill: str,
                              prefix: Optional[str], config: Config) -> None:
    """渲染一个 1×1 表格，用于引用块和结论箭头。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell = table.cell(0, 0)
    # 清空默认段落
    cell.text = ""
    para = cell.paragraphs[0]

    # 条件缩进：引用块不缩进（▌ 标记已提供引导），普通内容缩进
    if not config.typography.get("quote_no_indent", False):
        indent_cm = config.typography.get("first_line_indent_cm", 0.85)
        para.paragraph_format.first_line_indent = Cm(indent_cm)

    # 填充背景色
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{bg_fill}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)

    # 若无标记前缀，使用 _add_rich_text（内容里可能有 **粗体**）
    _add_rich_text(para, text, config)

    # 设置 cell 内所有文本为 body 字号
    for run in para.runs:
        run.font.size = Pt(config.typography.get("font_size", {}).get("body", 12))
        run.font.name = config.typography.get("font_family", {}).get("body", "Times New Roman")


# ── 富文本（粗体、箭头） ──────────────────────────────────

def _add_rich_text(para, text: str, config: Config,
                    segments: Optional[list[tuple[str, Optional[str]]]] = None) -> None:
    """添加带格式文本：支持 **粗体**、→/⇒ 前缀箭头和 [超链接](url)。

    Args:
        para: 段落对象。
        text: 原始 Markdown 文本（含 ** 标记）。
        config: 配置。
        segments: 可选，由 parser 提供的链接分段 [(文字, URL 或 None)]。
                   提供时优先使用此分段（URL 渲染为可点击超链接），
                   但文本中的 **粗体** 标记仍会解析。
    """
    # 渲染前压缩长破折号（仅视觉压缩，不改原文）
    text = _compress_dashes(text)

    arrow_m = _RE_ARROW.match(text)
    prefix = ""
    content = text
    if arrow_m:
        prefix = arrow_m.group(1) + " "
        content = arrow_m.group(2)

    if prefix:
        run = para.add_run(prefix)
        _set_font(run, "body", config)
        run.font.color.rgb = _ARROW_COLOR

    # 使用 segments（含链接标记）还是纯文本
    if segments:
        _add_rich_text_segments(para, content, segments, config)
    else:
        _add_rich_text_plain(para, content, config)


def _add_rich_text_plain(para, content: str, config: Config) -> None:
    """纯文本渲染（无超链接）。"""
    parts = _RE_BOLD.split(content)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        _set_font(run, "body", config)
        if i % 2 == 1:
            run.bold = True
        run.font.color.rgb = _BODY_COLOR


def _add_rich_text_segments(para, content: str,
                            segments: list[tuple[str, Optional[str]]],
                            config: Config) -> None:
    """带超链接的分段渲染。"""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement

    for seg_text, url in segments:
        if not seg_text:
            continue
        parts = _RE_BOLD.split(seg_text)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            _set_font(run, "body", config)
            if i % 2 == 1:
                run.bold = True
            run.font.color.rgb = _BODY_COLOR

            if url is not None:
                # 蓝色+下划线
                run.font.color.rgb = _hex_color("2056AE")
                run.underline = True
                # 创建超链接关系
                rel = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
                # 用 OxmlElement 创建 w:hyperlink
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), rel)
                # 把 run 移入 hyperlink
                run._element.addprevious(hyperlink)
                hyperlink.append(run._element)


def _add_cell_rich_text(para, cell_text: str, config: Config) -> None:
    """表格单元格富文本：支持 **粗体**，并把连续破折号压缩为单破折号。

    此前表格单元格直接 add_run 原始文本，导致 `**粗体**` 在 Word/PDF
    里泄漏为字面星号。这里与段落渲染器保持一致的分段粗体逻辑。
    """
    text = _compress_dashes(cell_text)
    parts = _RE_BOLD.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        _set_font(run, "body", config)
        if i % 2 == 1:
            run.bold = True
        run.font.color.rgb = _BODY_COLOR


# ── 多行表格 ───────────────────────────────────────────────

def _render_multi_row_table(doc: Document, block: Block, config: Config) -> None:
    rows = block.rows
    if not rows or len(rows) < 2:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_cell_rich_text(para, cell_text, config)
            for run in para.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
            if i == 0:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)


# ── 图表（DIAGRAM JSON） ──────────────────────────────────

def _render_diagram(doc: Document, block: Block, config: Config,
                    output_folder: Optional[str] = None,
                    diagram_counter: Optional[list] = None) -> Optional[dict]:
    """解析 DIAGRAM JSON 并生成图片，嵌入 Word + 保存到输出目录。"""
    data = block.diagram_data
    if not data:
        return None

    # focus 视角：若 DIAGRAM JSON 指定了 focus，裁剪为以该主体为中心的局部子图
    focus_id = data.get("focus")
    if focus_id:
        try:
            from viz_network import highlight_perspective
            focused = highlight_perspective(data, focus_id)
            if focused and focused.get("nodes"):
                data = focused
        except Exception:
            # focus 失败时回退到全景图，不阻断渲染
            pass

    title = data.get("title", "利益关系图")
    if diagram_counter is not None:
        diagram_counter[0] += 1
        seq = diagram_counter[0]
        seq_title = f"图{seq} {title}"
    else:
        seq = 0
        seq_title = title

    diag_result: Optional[dict] = None

    try:
        from viz_network import generate_diagram
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name

        result = generate_diagram(data, tmp_path)
        if result and os.path.exists(tmp_path):
            # 图片宽度按宽高比自适应
            img_w_cm = 14.0
            try:
                from PIL import Image
                with Image.open(tmp_path) as im:
                    w, h = im.size
                    ratio = w / h
                    if ratio > 1.8:
                        img_w_cm = 15.5
                    elif ratio < 1.0:
                        img_w_cm = 11.0
            except Exception:
                pass

            # 图片前留白
            doc.add_paragraph()
            # 插入图片
            doc.add_picture(tmp_path, width=Cm(img_w_cm))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 图题放在图下方（学术规范）
            if seq_title:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, seq_title, "body", config,
                         size_pt=10, color=_COVER_MUTED_COLOR)

            # 保存到输出目录
            if output_folder:
                safe_name = _safe_diagram_name(seq_title if seq else title)
                png_path = os.path.join(output_folder, f"{safe_name}.png")
                html_path = os.path.join(output_folder, f"{safe_name}.html")
                import shutil
                try:
                    shutil.copy2(tmp_path, png_path)
                except OSError:
                    pass  # 目标文件被锁时跳过，不影响报告生成
                try:
                    generate_diagram(data, html_path)
                except Exception:
                    pass
                diag_result = {"title": title, "seq": seq, "png": png_path, "html": html_path}

            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except ImportError:
        p = doc.add_paragraph()
        _add_run(p, f"[关系图: {title}]", "body", config,
                 size_pt=10, color=_COVER_MUTED_COLOR)
    except Exception:
        pass

    return diag_result


def _safe_diagram_name(title: str) -> str:
    """将图表标题转为安全文件名。"""
    safe = ""
    for ch in title:
        if ch.isalnum() or ch in "-_.":
            safe += ch
        elif ch in " ()（）【】":
            safe += "_"
        else:
            safe += "_"
    import re
    safe = re.sub(r"_+", "_", safe).strip("_")
    return safe if safe else "diagram"

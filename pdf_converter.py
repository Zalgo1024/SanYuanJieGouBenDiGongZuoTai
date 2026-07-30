"""
PDF 转换器 — 将 .docx 文件转换为 .pdf。

尝试以下方案（按优先级）：
1. LibreOffice --headless（跨平台，推荐）
2. pandoc + weasyprint/wkhtmltopdf
3. Windows COM（仅 Win32，需安装 Word）
4. 均不可用 → 返回空字符串并打印安装指引

环境变量:
    LIBREOFFICE_PATH  — 指定 soffice 安装路径，可传：
        * soffice.exe 完整路径
        * program/ 目录（如 C:\\Program Files\\LibreOffice\\program）
        * LibreOffice 根目录（如 C:\\Program Files\\LibreOffice）
      未设置时自动查 PATH 及 Windows 常见安装位置。
    ALLOW_WORD_COM_PDF — 设为 "1" 才启用 Word COM 转换（默认禁用）。

用法:
    from pdf_converter import convert_to_pdf
    result = convert_to_pdf("input.docx", "output.pdf")
"""

import os
import shutil
import subprocess
import sys
import tempfile
import time
import warnings


def _find_executable(name: str) -> bool:
    """检查可执行文件是否存在。"""
    if shutil.which(name) is not None:
        return True
    # Windows 回退：PATH 查不到时，直接尝试调用 --version
    if sys.platform == "win32":
        try:
            subprocess.run([name, "--version"], capture_output=True, timeout=5)
            return True
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            return False
    return False


def _resolve_soffice() -> str:
    """解析 soffice 可执行文件路径。

    优先级：
    1. ``LIBREOFFICE_PATH`` 环境变量 —— 可传 soffice 完整路径，
       或 ``program/`` 目录，或 LibreOffice 根目录均可。
    2. PATH 查找（``soffice`` / ``libreoffice``）
    3. Windows 常见安装位置（Program Files / LOCALAPPDATA）

    Returns:
        soffice 可执行文件的绝对路径；找不到返回空字符串。
    """
    exe_names = ("soffice.exe", "soffice", "libreoffice.exe", "libreoffice")

    # 1. LIBREOFFICE_PATH 环境变量
    env_val = os.environ.get("LIBREOFFICE_PATH", "").strip()
    if env_val:
        # 直接指向可执行文件
        if os.path.isfile(env_val):
            return os.path.abspath(env_val)
        # 指向目录：先在目录本身找
        for exe in exe_names:
            candidate = os.path.join(env_val, exe)
            if os.path.isfile(candidate):
                return os.path.abspath(candidate)
        # 指向 LibreOffice 根目录：尝试 program/ 子目录
        program_sub = os.path.join(env_val, "program")
        if os.path.isdir(program_sub):
            for exe in exe_names:
                candidate = os.path.join(program_sub, exe)
                if os.path.isfile(candidate):
                    return os.path.abspath(candidate)

    # 2. PATH 查找
    for name in exe_names:
        found = shutil.which(name)
        if found:
            return found

    # 3. Windows 常见安装位置
    if sys.platform == "win32":
        common_dirs = []
        pf = os.environ.get("PROGRAMFILES", "")
        if pf:
            common_dirs.append(os.path.join(pf, "LibreOffice", "program"))
        pf86 = os.environ.get("PROGRAMFILES(X86)", "")
        if pf86:
            common_dirs.append(os.path.join(pf86, "LibreOffice", "program"))
        la = os.environ.get("LOCALAPPDATA", "")
        if la:
            common_dirs.append(os.path.join(la, "Programs", "LibreOffice", "program"))
        for d in common_dirs:
            for exe in ("soffice.exe",):
                candidate = os.path.join(d, exe)
                if os.path.isfile(candidate):
                    return os.path.abspath(candidate)

    return ""


def _convert_libreoffice(docx_path: str, pdf_path: str) -> bool:
    """使用 LibreOffice 转换（保留超链接）。"""
    soffice = _resolve_soffice()
    if not soffice:
        return False
    # 使用 writer_pdf_Export 过滤器确保超链接保留
    cmd = [soffice, "--headless", "--convert-to", "pdf:writer_pdf_Export",
           "--outdir", os.path.dirname(pdf_path), docx_path]
    result = subprocess.run(cmd, capture_output=True, text=True,
                            timeout=120, errors="replace")
    if result.returncode == 0 and os.path.exists(pdf_path):
        return True
    # 回退：不带过滤器重试
    cmd2 = [soffice, "--headless", "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_path), docx_path]
    result2 = subprocess.run(cmd2, capture_output=True, text=True,
                             timeout=120, errors="replace")
    if result2.returncode == 0 and os.path.exists(pdf_path):
        return True
    # LibreOffice 自动命名输出文件（扩展名替换而非指定路径）
    # 注意：LibreOffice 可能忽略 --outdir，将 PDF 输出到源文件同目录。
    # 使用 copy2 而非 move，避免跨盘删除时触发沙箱 safe-delete 拦截。
    auto_pdf = os.path.splitext(docx_path)[0] + ".pdf"
    if os.path.exists(auto_pdf) and auto_pdf != pdf_path:
        shutil.copy2(auto_pdf, pdf_path)
        return True
    return False


def _extract_headings(docx_path: str):
    """从 docx 提取一级/二级标题（标题文本、层级），用于生成静态目录。

    跳过「目录」标题本身。返回 [(title, level), ...]，按文档出现顺序。
    """
    try:
        from docx import Document
    except ImportError:
        return []
    doc = Document(docx_path)
    headings = []
    for p in doc.paragraphs:
        style = (p.style.name if p.style is not None else "") or ""
        if style in ("Heading 1", "Heading 2"):
            t = p.text.strip()
            if t.replace(" ", "") == "目录":
                continue
            level = 1 if style == "Heading 1" else 2
            headings.append((t, level))
    return headings


def _measure_heading_pages(pdf_path: str, headings):
    """从 PDF 书签（outline）读取每个标题的真实页码（1-based）。

    书签由 docx 的正文标题（Heading 1/2）转换而来，页码精确，不受正文排版/对齐
    干扰；比逐页文本搜索可靠——文本搜索会误命中目录页本身，或被两端对齐插入的
    空格扰乱导致匹配失败/命中错误页。
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return [(t, l, 1) for (t, l) in headings]
    try:
        reader = PdfReader(pdf_path)
    except Exception:
        return [(t, l, 1) for (t, l) in headings]

    # 全角标点 → 半角，用于兜底匹配。防止 docx 标题与 LibreOffice 生成的 PDF
    # 书签在某些字符（全角冒号、引号、括号等）编码不一致时精确匹配失败、
    # 导致目录页码回退到 1。仅在精确匹配未命中时才启用，不影响正常匹配结果。
    _FW2HW = {
        '：': ':', '；': ';', '，': ',', '。': '.', '、': ',',
        '（': '(', '）': ')', '【': '[', '】': ']', '《': '<', '》': '>',
        '“': '"', '”': '"', '‘': "'", '’': "'", '·': '.',
        '—': '-', '－': '-', '～': '~', '！': '!', '？': '?',
    }

    def _norm_ws(s):
        """仅去空白（主匹配）。"""
        return "".join(str(s).split())

    def _norm_fuzzy(s):
        """去空白 + 全角标点转半角（兜底匹配）。"""
        s = "".join(str(s).split())
        return "".join(_FW2HW.get(ch, ch) for ch in s)

    # 扁平化书签 → [(去空白标题, 模糊标题, 页码1based), ...]
    bookmarks = []

    def _walk(items):
        for it in items:
            if isinstance(it, list):
                _walk(it)
                continue
            title = getattr(it, "title", None)
            if title is None:
                continue
            try:
                pn = reader.get_destination_page_number(it)
            except Exception:
                pn = None
            if pn is None:
                continue
            bookmarks.append((_norm_ws(title), _norm_fuzzy(title), int(pn) + 1))

    try:
        _walk(reader.outline)
    except Exception:
        pass

    if not bookmarks:
        return [(t, l, 1) for (t, l) in headings]

    result = []
    for title, level in headings:
        h_ws = _norm_ws(title)
        h_fz = _norm_fuzzy(title)
        page = 1
        # 主：精确（去空白）匹配
        for b_ws, b_fz, bpage in bookmarks:
            if b_ws == h_ws:
                page = bpage
                break
        else:
            # 兜底：全角/半角标点归一后匹配，避免页码回退到 1
            for b_ws, b_fz, bpage in bookmarks:
                if b_fz == h_fz:
                    page = bpage
                    break
        result.append((title, level, page))
    return result


def _inject_static_toc(docx_path: str, entries) -> bool:
    """把目录条目（标题、层级、页码）以「点引导符 + 页码」的格式写入 docx。

    定位「目录」标题段落与后续的分页符段落，清除两者之间的内容，
    再插入静态目录条目。目录标题与分页符保持不动。返回是否成功。
    """
    try:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Cm, Pt
    except ImportError:
        return False

    doc = Document(docx_path)

    # 定位标记段落
    toc_title_para = None
    pagebreak_para = None
    for p in doc.paragraphs:
        if toc_title_para is None and p.text.replace(" ", "") == "目录":
            toc_title_para = p
        if pagebreak_para is None and 'w:type="page"' in p._element.xml:
            pagebreak_para = p
    if toc_title_para is None or pagebreak_para is None:
        return False

    # 清除两者之间的所有段落
    body = doc.element.body
    nxt = toc_title_para._element.getnext()
    while nxt is not None and nxt is not pagebreak_para._element:
        to_remove = nxt
        nxt = nxt.getnext()
        body.remove(to_remove)

    # 计算右侧制表位（文本区右边界）
    sec = doc.sections[0]
    try:
        right_cm = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    except Exception:
        right_cm = 15.0
    if right_cm <= 0:
        right_cm = 15.0

    # 依次插入目录条目（插到「目录」标题之后）
    prev = toc_title_para._element
    for title, level, page in entries:
        para = doc.add_paragraph()
        # 右侧制表位 + 点引导符（用原生 OOXML，兼容各版本 python-docx）
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), str(int(right_cm * 567)))
        tabs.append(tab)
        para._p.get_or_add_pPr().append(tabs)
        if level >= 2:
            para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass
        tab = OxmlElement("w:tab")
        run._element.append(tab)
        run2 = para.add_run(str(page))
        run2.font.size = Pt(12)
        run2.font.name = "Times New Roman"
        try:
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass
        prev.addnext(para._element)
        prev = para._element

    doc.save(docx_path)
    return True


def _bake_toc(docx_path: str) -> bool:
    """将目录按真实页码写入 docx，使 LibreOffice 转换出的 PDF 也含完整目录。

    两步法避免目录自身长度影响页码测量：
      1) 先用占位页码注入目录（确定目录占用的篇幅）；
      2) 转一次临时 PDF，读每个标题的真实页码；
      3) 用真实页码重新注入目录。
    仅在 LibreOffice 可用时生效；任何一步失败都优雅回退（保留原 TOC 域）。
    """
    soffice = _resolve_soffice()
    if not soffice:
        return False
    headings = _extract_headings(docx_path)
    if not headings:
        return False

    # 第一步：占位页码注入
    if not _inject_static_toc(docx_path, [(t, l, 1) for (t, l) in headings]):
        return False

    # 第二步：测量真实页码
    import tempfile
    tmp_pdf = tempfile.mktemp(suffix=".pdf")
    measured = None
    try:
        if _convert_libreoffice(docx_path, tmp_pdf):
            measured = _measure_heading_pages(tmp_pdf, headings)
    finally:
        try:
            os.unlink(tmp_pdf)
        except Exception:
            pass

    if measured is None:
        # 测量失败：保留占位目录（页码=1，至少有目录可见）
        return True

    # 第三步：用真实页码重新注入
    _inject_static_toc(docx_path, measured)
    return True


def _convert_pandoc(docx_path: str, pdf_path: str) -> bool:
    """使用 pandoc 转换（需要安装 pdflatex 或 wkhtmltopdf）。"""
    # 尝试 pandoc → pdf (via pdflatex)
    cmd = ["pandoc", docx_path, "-o", pdf_path, "--from", "docx"]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0 and os.path.exists(pdf_path):
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return False


def _win32com_available() -> bool:
    """pywin32 是否可用（Word COM 转换的前置）。"""
    if sys.platform != "win32":
        return False
    try:
        import win32com.client  # noqa: F401
        return True
    except ImportError:
        return False


def _convert_win32_com(docx_path: str, pdf_path: str) -> bool:
    """使用 Windows COM (Word) 转换。仅限 Windows。

    健壮性要点（无头/自动化环境常见坑）：
    - gencache.EnsureDispatch：早期绑定，SaveAs/ExportAsFixedFormat 等方法名可被正确解析
      （延迟绑定会把 Documents.Open 的返回值误解析为方法本身，导致 AttributeError）。
    - DispatchEx：强制启动一个全新的 Word 实例，避免连接到已存在且卡在对话框的可见 Word。
    - DisplayAlerts=0 / ScreenUpdating=False：抑制「恢复文档 / 只读 / PDF 选项」等弹窗，
      否则 COM 呼叫会被拒绝（RPC_E_CALL_REJECTED）。
    - 对 RPC_E_CALL_REJECTED 重试：Word 启动首跑弹窗属瞬时，退避后通常可过。
    """
    if sys.platform != "win32":
        return False
    # 安全护栏：默认禁止在用户本机调用 Word COM 转 PDF。
    # 它会启动用户正在使用的 Word 实例，且异常退出时易留下僵尸进程、
    # 抢占 Word 的 COM 锁，导致用户无法正常打开/编辑 Word。
    # 仅在专用转换服务器上显式设置 ALLOW_WORD_COM_PDF=1 才启用。
    if os.environ.get("ALLOW_WORD_COM_PDF") != "1":
        warnings.warn("Word COM PDF 转换已禁用（默认）。如需启用，请设置 ALLOW_WORD_COM_PDF=1。")
        return False
    import pywintypes  # pywin32
    import win32com.client

    docx_abs = os.path.abspath(docx_path)
    pdf_abs = os.path.abspath(pdf_path)
    word = None
    # RPC_E_CALL_REJECTED = -2147418111
    for attempt in range(3):
        try:
            word = win32com.client.gencache.EnsureDispatch("Word.Application")
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            word.ScreenUpdating = False
            doc = word.Documents.Open(FileName=docx_abs, Visible=False)
            doc.SaveAs(pdf_abs, 17)  # 17 = wdFormatPDF
            doc.Close()
            return os.path.exists(pdf_abs)
        except pywintypes.com_error as e:
            code = e.args[0] if e.args else None
            if code == -2147418111 and attempt < 2:
                warnings.warn(f"Word COM 被拒绝（重试 {attempt + 1}）：{e}")
                time.sleep(2.0)
                continue
            warnings.warn(f"Word COM 转换失败：{e}")
            return False
        except ImportError:
            return False
        except Exception as e:  # noqa: BLE001
            warnings.warn(f"Word COM 转换失败：{e}")
            return False
        finally:
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
    return False


def diagnose_pdf() -> dict:
    """诊断当前环境可用的 PDF 转换器，避免「静默失败」。

    返回各转换器的可用性；recommended 给出推荐方案。
    """
    soffice_path = _resolve_soffice()
    return {
        "libreoffice": bool(soffice_path),
        "libreoffice_path": soffice_path,
        "libreoffice_env": os.environ.get("LIBREOFFICE_PATH", ""),
        "pandoc": bool(_find_executable("pandoc")),
        "word_com": _win32com_available(),
        "word_com_allowed": os.environ.get("ALLOW_WORD_COM_PDF") == "1",
        "recommended": "libreoffice",
        "reason": (
            "LibreOffice（无头、保留超链接，推荐）；其次 pandoc+LaTeX；"
            "Windows COM 兜底默认禁用，避免干扰用户本机 Word，"
            "仅在专用转换服务器设置 ALLOW_WORD_COM_PDF=1 时启用。"
            "可通过 LIBREOFFICE_PATH 环境变量指定 soffice 安装路径。"
        ),
    }


def convert_to_pdf(docx_path: str, pdf_path: str) -> str:
    """将 .docx 转换为 .pdf。

    Args:
        docx_path: 输入 .docx 文件路径（绝对或相对路径）。
        pdf_path: 输出 .pdf 文件路径。

    Returns:
        成功时返回 pdf_path，失败时返回空字符串。
    """
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(docx_path):
        warnings.warn(f"输入文件不存在: {docx_path}")
        return ""

    # 转换前：把目录按真实页码写回 docx，确保 LibreOffice 导出的 PDF 含完整目录。
    # 仅在 LibreOffice 可用时生效；失败则优雅回退（PDF 中可能无目录，但不报错）。
    if _resolve_soffice():
        try:
            _bake_toc(docx_path)
        except Exception as e:  # noqa: BLE001
            warnings.warn(f"目录预渲染失败（PDF 可能无目录）: {e}")

    # 按优先级尝试各方案
    converters = [
        ("LibreOffice", _convert_libreoffice),
        ("pandoc", _convert_pandoc),
    ]
    # Word COM 仅在显式允许时作为兜底（避免干扰用户本机 Word）
    if os.environ.get("ALLOW_WORD_COM_PDF") == "1":
        converters.append(("Windows COM (Word)", _convert_win32_com))

    for name, converter in converters:
        try:
            if converter(docx_path, pdf_path):
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    return pdf_path
        except Exception as e:
            warnings.warn(f"PDF 转换器 {name} 失败: {e}")
            continue

    # 全部失败：给出安装指引
    _print_install_guide()
    return ""


def _print_install_guide() -> None:
    """打印 PDF 转换器的安装指引。"""
    guide = """
╔══════════════════════════════════════════════════════════════╗
║  PDF 转换未完成：未找到可用的 PDF 转换工具                    ║
╠══════════════════════════════════════════════════════════════╣
║  安装以下任一工具即可启用 PDF 导出：                          ║
║                                                              ║
║  📄 LibreOffice（推荐，跨平台）                                ║
║     https://www.libreoffice.org/download/                     ║
║     安装后 soffice 会在 PATH 中自动发现                        ║
║     若不在 PATH，设置环境变量 LIBREOFFICE_PATH 指向安装目录：  ║
║       set LIBREOFFICE_PATH=C:\\Program Files\\LibreOffice\\program  ║
║     （可传 program/ 目录、soffice.exe 完整路径或根目录）       ║
║                                                              ║
║  📄 pandoc + MiKTeX（Windows）                                 ║
║     https://pandoc.org/installing.html                         ║
║     https://miktex.org/download                                ║
║                                                              ║
║  📄 Windows + Microsoft Word + pywin32                         ║
║     pip install pywin32  并设置 ALLOW_WORD_COM_PDF=1          ║
║                                                              ║
║  Word 文档已成功生成，可手动导出 PDF。                         ║
╚══════════════════════════════════════════════════════════════╝
"""
    warnings.warn(guide)

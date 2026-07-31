"""输入材料接口：粘贴长文本 / 上传 .txt .md .docx .pdf，以及列表/统计/详情/删除。"""
import io
import logging
import os
import uuid

from fastapi import APIRouter, File, Form, UploadFile
from pydantic import BaseModel

from app.db import SessionLocal
from app.models import Material

logger = logging.getLogger(__name__)

router = APIRouter()


class MaterialCreate(BaseModel):
    project_id: str | None = None
    title: str | None = None
    content_text: str = ""
    source_type: str = "paste"
    source: str | None = None  # 来源出处：链接 / 文号 / 出处说明
    tags: str | None = None  # 逗号分隔标签


def _read_text_bytes(raw: bytes) -> str:
    """按编码回退读取文本（中文环境常见 utf-8 / gbk）。"""
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def parse_uploaded_file(filename: str, raw: bytes) -> tuple[str, str, list[str]]:
    """解析上传文件，返回 (纯文本, source_type, warnings)。

    txt/md 直接解码；docx 用 python-docx 抽取段落+表格；
    pdf 用 pypdf 抽取每页文本，并检测「无文本 / 乱码 / 超大文件」三类告警。
    其它按文本兜底。warnings 为可读告警码列表（前端展示提示）。
    """
    lower = (filename or "").lower()
    # PDF 超大阈值：超过则拒绝解析，避免占用内存且通常无法抽取有效文本
    PDF_MAX_BYTES = 50 * 1024 * 1024  # 50MB（按用户要求调整）

    if lower.endswith(".txt"):
        return _read_text_bytes(raw), "txt", []
    if lower.endswith(".md"):
        return _read_text_bytes(raw), "md", []
    if lower.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
            warn = _text_quality_warnings(text, is_pdf=False)
            return text, "docx", warn
        except Exception as exc:  # noqa: BLE001
            logger.warning("docx 解析失败，按纯文本兜底：%s", exc)
            return _read_text_bytes(raw), "txt", []
    if lower.endswith(".pdf"):
        if len(raw) > PDF_MAX_BYTES:
            return (
                "",
                "pdf",
                [
                    "pdf_too_large",
                    f"文件过大（{len(raw)//(1024*1024)}MB），已超过 "
                    f"{PDF_MAX_BYTES//(1024*1024)}MB 上限，无法解析。请拆分后重新上传。",
                ],
            )
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            parts = [(pg.extract_text() or "") for pg in reader.pages]
            text = "\n".join(parts)
            warn = _text_quality_warnings(text, is_pdf=True)
            return text, "pdf", warn
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdf 解析失败，按纯文本兜底：%s", exc)
            return _read_text_bytes(raw), "txt", ["pdf_parse_failed"]
    return _read_text_bytes(raw), "txt", []


def _text_quality_warnings(text: str, is_pdf: bool) -> list[str]:
    """文本质量告警：PDF 无文本 / 乱码；返回告警码（可含说明）。空文本不告警非 PDF。"""
    warnings: list[str] = []
    stripped = text.strip()
    if is_pdf and len(stripped) < 20:
        warnings.append("pdf_text_empty")
    # 乱码检测：替换字符占比过高，或含大量 NUL/控制字符噪声
    if stripped:
        repl = stripped.count("�")
        ratio = repl / max(len(stripped), 1)
        nul = stripped.count("\x00")
        if ratio > 0.02 or nul > 5:
            warnings.append("pdf_garbled" if is_pdf else "text_garbled")
    return warnings


def _material_meta(m: Material) -> dict:
    return {
        "id": m.id,
        "project_id": m.project_id,
        "title": m.title,
        "source_type": m.source_type,
        "source": m.source,
        "tags": m.tags,
        "warnings": m.warnings or [],
        "original_filename": m.original_filename,
        "char_count": m.char_count,
        "created_at": m.created_at.isoformat() if m.created_at else None,
    }


def _material_full(m: Material) -> dict:
    d = _material_meta(m)
    d["content_text"] = m.content_text
    return d


@router.post("/api/materials")
def create_material(req: MaterialCreate):
    """手动粘贴长文本创建素材。可带 source（来源出处）与 tags（标签）。"""
    text = req.content_text or ""
    m = Material(
        id=uuid.uuid4().hex,
        project_id=req.project_id or None,
        title=(req.title or "未命名素材")[:200],
        content_text=text,
        source_type=req.source_type or "paste",
        source=req.source or None,
        tags=req.tags or None,
        char_count=len(text),
    )
    with SessionLocal() as db:
        db.add(m)
        db.commit()
        db.refresh(m)
        return _material_meta(m)


@router.post("/api/materials/upload")
async def upload_material(
    file: UploadFile = File(...),
    project_id: str = Form(None),
    title: str = Form(None),
    source: str = Form(None),
    tags: str = Form(None),
):
    """上传文件（.txt/.md/.docx/.pdf），按扩展名解析为纯文本。

    解析告警（PDF 无文本/乱码/超大）写入 warnings 并返回。
    """
    raw = await file.read()
    text, st, warn = parse_uploaded_file(file.filename or "file", raw)
    m = Material(
        id=uuid.uuid4().hex,
        project_id=project_id or None,
        title=(title or file.filename or "未命名素材")[:200],
        content_text=text,
        source_type=st,
        source=source or None,
        tags=tags or None,
        original_filename=file.filename,
        char_count=len(text),
        warnings=warn or None,
    )
    with SessionLocal() as db:
        db.add(m)
        db.commit()
        db.refresh(m)
        return _material_meta(m)


@router.get("/api/materials")
def list_materials(project_id: str | None = None, q: str | None = None):
    """素材列表：支持按项目过滤；q 关键词搜索（标题/正文/来源/标签）。"""
    with SessionLocal() as db:
        query = db.query(Material)
        if project_id:
            query = query.filter(Material.project_id == project_id)
        if q and q.strip():
            like = f"%{q.strip()}%"
            query = query.filter(
                (Material.title.ilike(like))
                | (Material.content_text.ilike(like))
                | (Material.source.ilike(like))
                | (Material.tags.ilike(like))
            )
        rows = query.order_by(Material.created_at.desc()).all()
        return [_material_meta(r) for r in rows]


@router.get("/api/materials/stats")
def materials_stats(project_id: str | None = None):
    """材料来源统计（数据页「材料来源统计」用，基于真实材料账本聚合）。

    - total：材料总数
    - by_type：按入库方式（paste/txt/md/docx/pdf）计数
    - by_source：按用户填写的来源出处（source 字段）聚合，取前 12 项
    - with_warnings：带解析告警（PDF 无文本/乱码/超大等）的材料数
    - linked_to_project：已关联到项目的材料数
    """
    with SessionLocal() as db:
        query = db.query(Material)
        if project_id:
            query = query.filter(Material.project_id == project_id)
        rows = query.all()
        by_type: dict[str, int] = {}
        by_source: dict[str, int] = {}
        with_warnings = 0
        linked = 0
        for m in rows:
            t = m.source_type or "unknown"
            by_type[t] = by_type.get(t, 0) + 1
            if m.source and m.source.strip():
                s = m.source.strip()
                by_source[s] = by_source.get(s, 0) + 1
            if m.warnings:
                with_warnings += 1
            if m.project_id:
                linked += 1
        top_sources = sorted(by_source.items(), key=lambda kv: kv[1], reverse=True)[:12]
        return {
            "total": len(rows),
            "by_type": by_type,
            "by_source": [{"source": s, "count": c} for s, c in top_sources],
            "with_warnings": with_warnings,
            "linked_to_project": linked,
        }


@router.get("/api/materials/{mid}")
def get_material(mid: str):
    with SessionLocal() as db:
        m = db.get(Material, mid)
        if not m:
            return {"error": "not_found"}
        return _material_full(m)


@router.delete("/api/materials/{mid}")
def delete_material(mid: str):
    with SessionLocal() as db:
        m = db.get(Material, mid)
        if not m:
            return {"error": "not_found"}
        db.delete(m)
        db.commit()
        return {"ok": True}

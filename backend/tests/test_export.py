"""Word / PDF 导出测试 —— 验证报告产物（docx / 关系图 / PDF 转换器诊断）真实可用。

测试聚焦：
- docx 能由域引擎真实生成，且是合法可读的 Word 文档（段落数、内容落盘）；
- 关系图 PNG / HTML 能真实写出且非空；
- PDF 转换器诊断接口返回预期字段；
- convert_to_pdf 按环境行为正确：装了 LibreOffice 则真实转换成功；
  无可用转换器则优雅返回空串（不抛异常、不阻断主流程）。

域引擎（engine / viz_network / pdf_converter）位于 ENGINE_DIR，不属于本仓库源码。
为避免「缺依赖 / 未配置 ENGINE_DIR」时整套 pytest 收集即崩溃，这里把引擎导入延迟到
fixture，并在不可用时优雅 skip 相关用例 —— 这样 `pytest tests/` 在任何环境下都能完整
收集与运行，无需 `--ignore=tests/test_export.py`。
"""
import json
import os
import pytest
from pathlib import Path

import docx

from app import rule_engine
from app.engine_bridge import export_report, _ensure_engine_on_path

FIX = Path(__file__).resolve().parent / "fixtures"


def _event_md() -> str:
    data = json.loads((FIX / "sample_event.json").read_text(encoding="utf-8"))
    return rule_engine.generate(rule_engine.StructuredInput.model_validate(data))


@pytest.fixture(scope="module")
def engine_modules():
    """域引擎（engine / viz_network / pdf_converter）位于 engine_dir，可用才跑导出测试。

    引擎不是本仓库依赖，导入失败（缺 matplotlib/networkx 或未配置 ENGINE_DIR）时，
    相关用例优雅 skip，而不是让整套收集的导入阶段就崩溃。
    """
    _ensure_engine_on_path()
    try:
        import engine  # noqa: F401
        import viz_network  # noqa: F401
        import pdf_converter  # noqa: F401
    except Exception as e:  # pragma: no cover - 依赖环境
        pytest.skip(f"域引擎不可导入（缺依赖或 ENGINE_DIR 未配置），跳过导出测试：{e}")
    return engine, viz_network, pdf_converter


def test_docx_export(tmp_path, engine_modules):
    md = _event_md()
    out = export_report("导出测试报告", md, output_dir=str(tmp_path), slug="export1")
    word = out.get("word")
    assert word and os.path.exists(word), f"docx 未生成: {out}"

    d = docx.Document(word)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    assert len(paras) > 5, "docx 正文段落过少"
    blob = "\n".join(paras)
    assert "利益" in blob, "报告内容应落盘到 docx"

    diagrams = out.get("diagrams") or []
    assert diagrams, "缺少关系图产物"
    first = diagrams[0]
    assert os.path.exists(first["png"]) or os.path.exists(first["html"]), "关系图文件未写出"


def test_viz_generate_both(tmp_path, engine_modules):
    _engine, viz_network, _pdf = engine_modules
    data = {
        "viz": "network",
        "title": "t",
        "nodes": [
            {"id": "a", "label": "A", "type": "actor"},
            {"id": "b", "label": "B", "type": "event"},
        ],
        "edges": [{"source": "a", "target": "b", "label": "x", "type": "economic"}],
    }
    png = str(tmp_path / "g.png")
    html = str(tmp_path / "g.html")
    r_png, r_html = viz_network.generate_both(data, png, html)
    assert r_png and os.path.exists(r_png) and os.path.getsize(r_png) > 0
    assert r_html and os.path.exists(r_html) and os.path.getsize(r_html) > 0


def test_pdf_diagnose_shape(engine_modules):
    _engine, _viz, pdf_converter = engine_modules
    diag = pdf_converter.diagnose_pdf()
    for k in ("libreoffice", "pandoc", "word_com", "recommended"):
        assert k in diag, f"diagnose_pdf 缺字段 {k}"
    assert isinstance(diag["libreoffice"], bool)
    assert isinstance(diag["word_com"], bool)


def test_pdf_convert_graceful_when_no_converter(tmp_path, engine_modules):
    _engine, _viz, pdf_converter = engine_modules
    p = tmp_path / "mini.docx"
    docx.Document().save(str(p))
    pdf = str(tmp_path / "mini.pdf")
    # 行为按环境自适应：装了 LibreOffice → 真实转换成功；无转换器 → 优雅返回空串
    diag = pdf_converter.diagnose_pdf()
    res = pdf_converter.convert_to_pdf(str(p), pdf)
    if diag.get("libreoffice"):
        assert res and os.path.exists(res), f"本机已装 LibreOffice，PDF 转换应成功，却得到: {res}"
    else:
        assert res == ""
        assert not os.path.exists(pdf)
